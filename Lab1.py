
import docx

table_manage = {
    "01000": "return",
    "00010": "slash_n",
    "11111": "eng",
    "11011": "nums",
    "00100": "space",
    "00000": "rus"
}

any_alphabet = {
    "00011": ['A', 'А', '-'],
    "11001": ['B', 'Б', '?'],
    "01110": ['C', 'Ц', ':'],
    "01001": ['D', 'Д', 'Кто там?'],
    "00001": ['E', 'Е', '3'],
    "01101": ['F', 'Ф', 'Э'],
    "11010": ['G', 'Г', 'Ш'],
    "10100": ['H', 'Х', 'Щ'],
    "00110": ['I', 'И', '8'],
    "01011": ['J', 'Й', 'Ю'],
    "01111": ['K', 'К', '('],
    "10010": ['L', 'Л', ')'],
    "11100": ['M', 'М', '.'],
    "01100": ['N', 'Н', ','],
    "11000": ['O', 'О', '9'],
    "10110": ['P', 'П', '0'],
    "10111": ['Q', 'Я', '1'],
    "01010": ['R', 'Р', '4'],
    "00101": ['S', 'С', "'"],
    "10000": ['T', 'Т', '5'],
    "00111": ['U', 'У', '7'],
    "11110": ['V', 'Ж', '='],
    "10011": ['W', 'В', '2'],
    "11101": ['X', 'Ь', '/'],
    "10101": ['Y', 'Ы', '6'],
    "10001": ['Z', 'З', '+']
}

# w:color
def get_character_font_color(docx_path):
    doc = docx.Document(docx_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            xml_element = run._element
            if 'w:color' in xml_element.xml:
                # Извлекаем информацию о цвете шрифта
                font_color_info = xml_element.find('.//w:color', namespaces=xml_element.nsmap)
                if font_color_info is not None:
                    return font_color_info.attrib, 'w:color'

    return None, None
# w:highlight
def get_character_highlight(docx_path):
    doc = docx.Document(docx_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            xml_element = run._element
            if 'w:highlight' in xml_element.xml:
                # Извлекаем информацию о фоне символов
                char_highlight_info = xml_element.find('.//w:highlight', namespaces=xml_element.nsmap)
                if char_highlight_info is not None:
                    return char_highlight_info.attrib, 'w:highlight'

    return None, None
# w:sz
def get_character_font_size(docx_path):
    doc = docx.Document(docx_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            xml_element = run._element
            if 'w:sz' in xml_element.xml:
                # Извлекаем информацию о размере шрифта
                font_size_info = xml_element.find('.//w:sz', namespaces=xml_element.nsmap)
                if font_size_info is not None:
                    return font_size_info.attrib, 'w:sz'

    return None, None
# w:w
def get_character_font_scaling(docx_path):
    doc = docx.Document(docx_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            xml_element = run._element
            if 'w:w' in xml_element.xml:
                # Извлекаем информацию о масштабе шрифта
                font_scale_info = xml_element.find('.//w:w', namespaces=xml_element.nsmap)
                if font_scale_info is not None:
                    return font_scale_info.attrib, 'w:w'

    return None, None
# w:spacing
def get_character_spacing(docx_path):
    doc = docx.Document(docx_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            xml_element = run._element
            if 'w:spacing' in xml_element.xml:
                # Извлекаем информацию о межсимвольном интервале
                spacing_info = xml_element.find('.//w:spacing', namespaces=xml_element.nsmap)
                if spacing_info is not None:
                    return spacing_info.attrib, 'w:spacing'

    return None, None

def decode_baudot(array):
    strange_line = ''
    flag = 0
    for i in array:
        if i in table_manage:
            manager = table_manage[i]
            if manager == 'eng':
                flag = 0
            elif manager == 'rus':
                flag = 1
            elif manager == 'nums':
                flag = 2
            elif manager == '(return)':
                strange_line += '(CR)'
            elif manager == '(slash_n)':
                strange_line += '(SN)'
            elif manager == '(space)':
                strange_line += ' '

        elif i in any_alphabet:
            strange_line += ''.join(any_alphabet[i][flag])

    return strange_line






docx_path = 'variant04.docx'






# Информация об изменении цвета шрифта
character_font_color, color_flag = get_character_font_color(docx_path)
try:
    q = str(character_font_color).split(':')[2][1:-1]
except:
    1
if (character_font_color is not None) and (q != "'000000'"):
    print(f"Информация об изменении цвета шрифта: {q}")
elif (character_font_color is None) or (q == "'000000'"):
    print("Изменения цвета шрифта не найдены.")
    color_flag = None

# Информация о цвете бэка
character_highlight, highlight_flag = get_character_highlight(docx_path)
try:
    q = str(character_highlight).split(':')[2][1:-1]
except:
    1
if (character_highlight is not None) and (q != "'white'"):
    print(f"Информация об изменении цвета бэка: {q}")
elif (character_highlight is None) or (q == "'white'"):
    print(f"Изменения цвета бэка не найдены")
    highlight_flag = None

# Информация о размере шрифта
character_font_size, size_flag = get_character_font_size(docx_path)
try:
    q = str(character_font_size).split(':')[2][1:-1]
except:
    1
if (character_font_size is not None) and (q != "'24'"):
    print(f"Информация об изменении размера шрифта: {q}")
elif (character_font_size is None) or (q == "'24'"):
    print("Изменения размера шрифта не найдены.")
    size_flag = None

# Информация о масштабе шрифта
character_scaling, scale_flag = get_character_font_scaling(docx_path)
try:
    q = str(character_scaling).split(':')[2][:-1]
except:
    1
if (character_scaling is not None):
    print(f"Информация об изменении масштаба шрифта: {q}")
elif (character_scaling is None):
    print("Изменения масштаба шрифта не найдены.")
    scale_flag = None

# Информация о межсимвольном интервале
character_spacing, space_flag = get_character_spacing(docx_path)
try:
    q = str(character_spacing).split(':')[2][:-1]
except:
    1
if character_spacing is not None:
    print(f"Информация об изменениях межсимвольного интервала: {q}")
elif character_spacing is None:
    print("Изменения межсимвольного интервала не найдены.")
    space_flag = None

print('')

# 0's / 1's

viewDoc = docx.Document(docx_path)

flags_array = [color_flag, highlight_flag, size_flag, scale_flag, space_flag]
for i in flags_array:
    if i:
        parameter = i
        break
print(parameter)
# print(flags_array)
print()
ones_zeros_line = ''

for paragraph in viewDoc.paragraphs:
    for i in range(len(paragraph.runs)):
        print(paragraph.runs[i].text)
        if parameter in paragraph.runs[i]._element.xml:
            char_info = paragraph.runs[i]._element.find(f'.//{parameter}', namespaces=paragraph.runs[i]._element.nsmap)
            if char_info is not None:
                char_info_q = char_info.attrib
                char_info_qw = str(char_info_q).split(':')[2][2:-2]
                print(f'{parameter} = {char_info_qw} \n')
                #print(len(str(paragraph.runs[i].text)))

                if char_info_qw == '010101' and color_flag:
                    ones_zeros_line += '1' * len(str(paragraph.runs[i].text))
                    #print('color ones')
                elif char_info_qw == '000000' and color_flag:
                    ones_zeros_line += '0' * len(str(paragraph.runs[i].text))
                    #print('color zeros')

                if char_info_qw != 'white' and highlight_flag:
                    ones_zeros_line += '1' * len(str(paragraph.runs[i].text))
                    #print('HL ones')

                if char_info_qw == '23' and size_flag:
                    ones_zeros_line += '1' * len(str(paragraph.runs[i].text))
                    #print('size ones')

                elif char_info_qw == '24' and size_flag:
                    ones_zeros_line += '0' * len(str(paragraph.runs[i].text))
                    #print('size zeros')

                if char_info_qw == '99' and scale_flag:
                    ones_zeros_line += '1' * len(str(paragraph.runs[i].text))
                    #print('scale ones')

                if char_info_qw == '20' and space_flag:
                    ones_zeros_line += '1' * len(str(paragraph.runs[i].text))
                    #print('space ones')

        elif parameter not in paragraph.runs[i]._element.xml:
            print(f'{parameter} = без изменений \n')
            ones_zeros_line += '0' * len(str(paragraph.runs[i].text))

print(ones_zeros_line)

ones_zeros_line_to_five = []
for i in range(len(ones_zeros_line)):
    try:
        ones_zeros_line[i+5]
        temp_line = ''
        for j in range(5):
            temp_line += ones_zeros_line[i+j]
        i += 5
        if len(temp_line) == 5:
            ones_zeros_line_to_five.append(temp_line)

    except:
        break

decoded_five_line = decode_baudot(ones_zeros_line_to_five)
print(ones_zeros_line_to_five)
print(f'Baudot decoding: {decoded_five_line}')

ones_zeros_line_to_byte = bytearray(int(ones_zeros_line[i:i+8], 2) for i in range(0, len(ones_zeros_line), 8))

print(f"KOI-8R decoding: {ones_zeros_line_to_byte.decode('koi8-r')}")
print(f"cp866 decoding: {ones_zeros_line_to_byte.decode('cp866')}")
print(f"windows1251: {ones_zeros_line_to_byte.decode('windows-1251')}")

